"""Build .docx files from the manuscript markdown sources.

Reads:  /Users/alexvintera/Desktop/crc-metagenomics/manuscript/markdown/*.md
Writes:
  - manuscript_complete.md  (concatenated)
  - CRC_Title_Page.docx
  - CRC_Abstract.docx
  - CRC_Introduction.docx
  - CRC_Methods.docx
  - CRC_Results.docx
  - CRC_Discussion.docx
  - CRC_References.docx
  - CRC_Table1.docx
  - Supplementary_Tables.docx
  - CRC_Manuscript_Complete.docx
"""

import os
import re
from docx import Document
from docx.shared import Pt

MD_DIR = "/Users/alexvintera/Desktop/crc-metagenomics/manuscript/markdown"
OUT_DIR = "/Users/alexvintera/Desktop/crc-metagenomics/manuscript"

SECTIONS = {
    "00_title.md": "CRC_Title_Page.docx",
    "01_abstract.md": "CRC_Abstract.docx",
    "02_introduction.md": "CRC_Introduction.docx",
    "03_methods.md": "CRC_Methods.docx",
    "04_results.md": "CRC_Results.docx",
    "05_discussion.md": "CRC_Discussion.docx",
    "06_references.md": "CRC_References.docx",
}


def read_md(name):
    with open(os.path.join(MD_DIR, name), "r", encoding="utf-8") as f:
        return f.read()


# Private-use placeholders for escaped markers.
ESC_AST = ""
ESC_CAR = ""
ESC_BTK = ""
ESC_USC = ""


def add_runs_with_inline_formatting(para, text):
    text = (text
            .replace(r"\*", ESC_AST)
            .replace(r"\^", ESC_CAR)
            .replace(r"\`", ESC_BTK)
            .replace(r"\_", ESC_USC))

    pattern = re.compile(
        r"(\*\*[^*]+\*\*)"
        r"|(\*[^*]+\*)"
        r"|(`[^`]+`)"
        r"|(\^[^\s\^]+\^)"
    )

    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            _add_run(para, text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            _add_run(para, tok[2:-2], bold=True)
        elif tok.startswith("*"):
            _add_run(para, tok[1:-1], italic=True)
        elif tok.startswith("`"):
            _add_run(para, tok[1:-1], mono=True)
        elif tok.startswith("^"):
            _add_run(para, tok[1:-1], superscript=True)
        pos = m.end()
    if pos < len(text):
        _add_run(para, text[pos:])


def _add_run(para, text, bold=False, italic=False, mono=False, superscript=False):
    text = (text
            .replace(ESC_AST, "*")
            .replace(ESC_CAR, "^")
            .replace(ESC_BTK, "`")
            .replace(ESC_USC, "_"))
    r = para.add_run(text)
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if mono:
        r.font.name = "Courier New"
    if superscript:
        r.font.superscript = True


def parse_table(lines, start_idx):
    header = lines[start_idx]
    if start_idx + 1 >= len(lines) or "---" not in lines[start_idx + 1]:
        return None, start_idx + 1
    rows = [header]
    i = start_idx + 2
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append(lines[i])
        i += 1

    def split_row(r):
        parts = [c.strip() for c in r.strip().strip("|").split("|")]
        return parts

    data = [split_row(r) for r in rows]
    return data, i


def add_table(doc, data):
    n_rows = len(data)
    n_cols = max(len(r) for r in data)
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    try:
        tbl.style = "Light Grid Accent 1"
    except KeyError:
        tbl.style = "Table Grid"
    for ri, row in enumerate(data):
        for ci in range(n_cols):
            cell = tbl.rows[ri].cells[ci]
            text = row[ci] if ci < len(row) else ""
            cell.text = ""
            para = cell.paragraphs[0]
            add_runs_with_inline_formatting(para, text)
            if ri == 0:
                for run in para.runs:
                    run.bold = True


def render_markdown_to_doc(doc, md):
    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            p = doc.add_heading(level=min(level, 4))
            add_runs_with_inline_formatting(p, text)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and "---" in lines[i + 1]:
            data, ni = parse_table(lines, i)
            if data:
                add_table(doc, data)
                i = ni
                continue

        if re.match(r"^[-*]\s+", stripped):
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
                item = re.sub(r"^[-*]\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Bullet")
                add_runs_with_inline_formatting(p, item)
                i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                item = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Number")
                add_runs_with_inline_formatting(p, item)
                i += 1
            continue

        buf = [line]
        i += 1
        while i < len(lines) and lines[i].strip():
            nxt = lines[i].strip()
            if (nxt.startswith("#")
                or nxt == "---"
                or nxt.startswith("|")
                or re.match(r"^[-*]\s+", nxt)
                or re.match(r"^\d+\.\s+", nxt)):
                break
            buf.append(lines[i])
            i += 1
        para_text = " ".join(s.strip() for s in buf)
        p = doc.add_paragraph()
        add_runs_with_inline_formatting(p, para_text)


def set_default_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def build_section(md_name, docx_name):
    md = read_md(md_name)
    doc = Document()
    set_default_style(doc)
    render_markdown_to_doc(doc, md)
    out_path = os.path.join(OUT_DIR, docx_name)
    doc.save(out_path)
    print(f"wrote {out_path}")


def build_table1_and_supplementary():
    md = read_md("07_supplementary.md")
    parts = re.split(r"\n---\n", md, maxsplit=1)
    table1_md = parts[0]
    supp_md = parts[1] if len(parts) > 1 else ""

    t1_doc = Document()
    set_default_style(t1_doc)
    # For the standalone Table 1 docx, drop the combined-section H1 entirely;
    # the table's own subheading ("Table 1. Demographic ...") is sufficient.
    table1_md_standalone = re.sub(
        r"^# Table 1 and Supplementary Material\s*\n+",
        "",
        table1_md,
        count=1,
    )
    render_markdown_to_doc(t1_doc, table1_md_standalone)
    t1_path = os.path.join(OUT_DIR, "CRC_Table1.docx")
    t1_doc.save(t1_path)
    print(f"wrote {t1_path}")

    supp_doc = Document()
    set_default_style(supp_doc)
    supp_md_full = "# Supplementary Material\n\n" + supp_md
    render_markdown_to_doc(supp_doc, supp_md_full)
    supp_path = os.path.join(OUT_DIR, "Supplementary_Tables.docx")
    supp_doc.save(supp_path)
    print(f"wrote {supp_path}")


def build_complete():
    title = read_md("00_title.md")
    abstract = read_md("01_abstract.md")
    intro = read_md("02_introduction.md")
    methods = read_md("03_methods.md")
    results = read_md("04_results.md")
    discussion = read_md("05_discussion.md")
    refs = read_md("06_references.md")
    supp = read_md("07_supplementary.md")

    combined = "\n\n---\n\n".join([title, abstract, intro, methods, results,
                                   discussion, refs, supp])
    out_md = os.path.join(MD_DIR, "manuscript_complete.md")
    with open(out_md, "w", encoding="utf-8") as f:
        f.write(combined)
    print(f"wrote {out_md}")

    doc = Document()
    set_default_style(doc)
    render_markdown_to_doc(doc, combined)
    out_docx = os.path.join(OUT_DIR, "CRC_Manuscript_Complete.docx")
    doc.save(out_docx)
    print(f"wrote {out_docx}")


def main():
    for md_name, docx_name in SECTIONS.items():
        build_section(md_name, docx_name)
    build_table1_and_supplementary()
    build_complete()


if __name__ == "__main__":
    main()
