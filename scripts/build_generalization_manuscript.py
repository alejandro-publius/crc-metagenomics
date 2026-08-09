#!/usr/bin/env python3
"""Build the coauthor-review DOCX for the generalization-risk manuscript."""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript/generalization_risk/manuscript.md"
OUTPUT = ROOT / "manuscript/generalization_risk/Generalization_Risk_Manuscript.docx"
FIGURES = ROOT / "manuscript/generalization_risk/figures"

# narrative_proposal preset: Letter, 1-inch margins, Calibri 11, justified,
# 8 pt paragraph spacing, 1.333 lines, restrained blue heading ladder.
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x55, 0x55)


def set_font(run, size: float = 11, bold=None, italic=None, color=None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.extend([begin, text, end])


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("mSystems Research Article | Coauthor review draft"), 8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("Page "), 9, color=MUTED)
    add_field(footer, "PAGE")


def add_inline(paragraph, text: str) -> None:
    """Render a small Markdown subset used by the manuscript."""
    text = re.sub(r"\^([12])\^", r"\1", text)
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`|\*[^*]+?\*)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            set_font(paragraph.add_run(text[position:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            set_font(paragraph.add_run(token[2:-2]), bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Courier New")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Courier New")
            run.font.size = Pt(9.5)
        else:
            set_font(paragraph.add_run(token[1:-1]), italic=True)
        position = match.end()
    if position < len(text):
        set_font(paragraph.add_run(text[position:]))


def add_title_page(doc: Document, lines: list[str]) -> int:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(36)
    title.paragraph_format.space_after = Pt(18)
    set_font(title.add_run(lines[0][2:]), 22, bold=True, color=DARK_BLUE)

    for raw in lines[1:7]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        add_inline(p, raw.replace("  ", ""))
    doc.add_page_break()
    return 7


def add_figure(doc: Document, filename: str, caption: str) -> None:
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(FIGURES / filename), width=Inches(6.35))
    shape._inline.docPr.set("descr", caption)
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.LEFT
    c.paragraph_format.keep_with_next = False
    set_font(c.add_run(caption), 9.5, italic=True, color=MUTED)


def main() -> None:
    lines = SOURCE.read_text().splitlines()
    doc = Document()
    style_document(doc)
    index = add_title_page(doc, lines)
    buffer: list[str] = []

    def flush() -> None:
        nonlocal buffer
        if buffer:
            p = doc.add_paragraph()
            add_inline(p, " ".join(part.strip() for part in buffer))
            buffer = []

    for raw in lines[index:]:
        line = raw.strip()
        if not line:
            flush()
            continue
        if line.startswith("### "):
            flush(); doc.add_heading(line[4:], level=2)
        elif line.startswith("## "):
            flush(); doc.add_heading(line[3:], level=1)
        elif re.match(r"^\d+\. ", line):
            flush()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
            add_inline(p, line)
        else:
            buffer.append(line)
    flush()

    add_figure(
        doc, "PortabilityLandscape.png",
        "Figure 1. Portability across biological representations. Points are held-out development cohorts; the red star is the untouched external cohort.",
    )
    add_figure(
        doc, "GeneralizationRisk.png",
        "Figure 2. Historical and label-free estimates versus observed target AUC. The red star is the external species model.",
    )
    doc.core_properties.title = "Biological detail does not guarantee portability of colorectal cancer metagenomic classifiers"
    doc.core_properties.subject = "Coauthor review draft"
    doc.core_properties.author = "Alejandro Velazquez; Rachel Selbrede"
    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
